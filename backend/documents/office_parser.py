"""
Parser per documenti Office (DOCX e DOC)
Converte DOCX/DOC ↔ HTML per editing real-time
Supporta .doc tramite conversione automatica con LibreOffice
"""
import os
import tempfile
import subprocess
import shutil
from pathlib import Path
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from bs4 import BeautifulSoup
import html2text


class OfficeDocumentParser:
    """
    Parser per documenti Office con conversione bidirezionale DOCX ↔ HTML
    Supporta .doc (Word 97-2003) tramite conversione automatica
    """
    
    @staticmethod
    def _find_libreoffice() -> Optional[str]:
        """
        Trova l'eseguibile di LibreOffice/OpenOffice nel sistema
        
        Returns:
            Path dell'eseguibile o None se non trovato
        """
        # Percorsi comuni per LibreOffice/OpenOffice
        possible_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
            '/usr/bin/libreoffice',  # Linux
            '/usr/bin/soffice',  # Linux alternative
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',  # Windows 32-bit
        ]
        
        # Prova a trovare nel PATH
        soffice = shutil.which('soffice') or shutil.which('libreoffice')
        if soffice:
            return soffice
        
        # Cerca nei percorsi comuni
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    @staticmethod
    def _convert_doc_to_docx(doc_path: str) -> Optional[str]:
        """
        Converte un file .doc in .docx usando LibreOffice headless
        
        Args:
            doc_path: Percorso del file .doc
            
        Returns:
            Percorso del file .docx convertito o None se errore
        """
        soffice = OfficeDocumentParser._find_libreoffice()
        
        if not soffice:
            return None
        
        # Crea directory temporanea per conversione
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Converti .doc → .docx usando LibreOffice headless
            # --headless: modalità senza GUI
            # --convert-to docx: formato di output
            # --outdir: directory di output
            cmd = [
                soffice,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                doc_path
            ]
            
            # Esegui conversione (timeout 30 secondi)
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                print(f"❌ LibreOffice conversion error: {result.stderr}")
                return None
            
            # Trova il file .docx generato
            doc_filename = os.path.basename(doc_path)
            docx_filename = os.path.splitext(doc_filename)[0] + '.docx'
            docx_path = os.path.join(temp_dir, docx_filename)
            
            if os.path.exists(docx_path):
                return docx_path
            else:
                return None
                
        except subprocess.TimeoutExpired:
            print("❌ LibreOffice conversion timeout")
            return None
        except Exception as e:
            print(f"❌ Errore conversione .doc: {e}")
            return None
    
    @staticmethod
    def docx_to_html(file_path: str) -> Dict[str, any]:
        """
        Converte un file DOCX in HTML editabile
        
        Args:
            file_path: Percorso del file DOCX
            
        Returns:
            Dict con 'html' (contenuto), 'metadata' (info documento), 'success' (bool)
        """
        if not DOCX_AVAILABLE:
            return {
                'html': '<p>⚠️ python-docx non installato. Esegui: pip install python-docx</p>',
                'metadata': {},
                'success': False,
                'error': 'python-docx non installato'
            }
        
        # Verifica estensione file
        file_ext = os.path.splitext(file_path)[1].lower()
        original_file_path = file_path
        converted_docx_path = None
        temp_dir_to_cleanup = None
        
        # Se è .doc, prova a convertirlo in .docx
        if file_ext == '.doc':
            print(f"📄 Rilevato file .doc, tento conversione automatica...")
            converted_docx_path = OfficeDocumentParser._convert_doc_to_docx(file_path)
            
            if converted_docx_path:
                print(f"✅ Conversione .doc → .docx completata")
                file_path = converted_docx_path
                temp_dir_to_cleanup = os.path.dirname(converted_docx_path)
                file_ext = '.docx'
            else:
                # Conversione fallita - mostra messaggio informativo
                return {
                    'html': '''
                        <div style="padding: 20px; text-align: center; font-family: Arial, sans-serif;">
                            <h2 style="color: #e67e22;">⚠️ LibreOffice Non Disponibile</h2>
                            <p style="margin: 20px 0;">Il formato <strong>.DOC</strong> (Word 97-2003) richiede LibreOffice per la conversione automatica.</p>
                            <p style="margin: 20px 0; color: #7f8c8d;">
                                <strong>Opzione 1:</strong> Installa LibreOffice sul server<br/>
                                <code>brew install libreoffice</code> (macOS)<br/>
                                <code>sudo apt-get install libreoffice</code> (Linux)
                            </p>
                            <p style="margin: 20px 0; color: #7f8c8d;">
                                <strong>Opzione 2:</strong> Converti manualmente in .DOCX
                            </p>
                            <ol style="text-align: left; display: inline-block; margin: 20px auto; color: #34495e;">
                                <li>Apri il file in Word/LibreOffice</li>
                                <li>Vai su <strong>File → Salva con nome</strong></li>
                                <li>Seleziona formato <strong>Word Document (.docx)</strong></li>
                                <li>Carica nuovamente il file nelle impostazioni</li>
                            </ol>
                        </div>
                    ''',
                    'metadata': {
                        'unsupported_format': True,
                        'extension': file_ext,
                        'conversion_failed': True
                    },
                    'success': False,
                    'error': 'LibreOffice non disponibile per conversione .doc'
                }
        
        # Verifica che ora sia .docx
        if file_ext not in ['.docx']:
            return {
                'html': f'<p>⚠️ Formato file non supportato: {file_ext}. Solo .DOC e .DOCX sono supportati.</p>',
                'metadata': {},
                'success': False,
                'error': f'Formato {file_ext} non supportato'
            }
        
        try:
            doc = Document(file_path)
            
            # Estrai metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
            }
            
            # Converti contenuto in HTML (stile OnlyOffice)
            html_parts = []
            
            for para in doc.paragraphs:
                # Salta paragrafi vuoti
                if not para.text.strip():
                    html_parts.append('<p><br/></p>')
                    continue
                
                # Estrai stile e formattazione del paragrafo
                style = para.style.name if para.style else 'Normal'
                alignment = para.alignment
                
                # Determina allineamento CSS
                align_css = ''
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align_css = 'text-align: center;'
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align_css = 'text-align: right;'
                elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    align_css = 'text-align: justify;'
                else:
                    align_css = 'text-align: left;'
                
                # Determina tag e stili basato sullo stile Word
                if 'Heading 1' in style or 'Title' in style:
                    tag = 'h1'
                    base_style = f'{align_css} font-size: 20pt; font-weight: bold; margin-top: 12pt; margin-bottom: 6pt;'
                elif 'Heading 2' in style:
                    tag = 'h2'
                    base_style = f'{align_css} font-size: 16pt; font-weight: bold; margin-top: 10pt; margin-bottom: 4pt;'
                elif 'Heading 3' in style:
                    tag = 'h3'
                    base_style = f'{align_css} font-size: 14pt; font-weight: bold; margin-top: 8pt; margin-bottom: 4pt;'
                else:
                    tag = 'p'
                    base_style = f'{align_css} font-size: 14pt; margin-top: 0pt; margin-bottom: 8pt; line-height: 1.3;'
                
                # Costruisci paragrafo HTML con formattazione inline
                para_html = f'<{tag} style="{base_style}">'
                
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    
                    # Escape HTML
                    import html as html_module
                    text = html_module.escape(text)
                    
                    # Costruisci stili inline per il run
                    run_styles = []
                    
                    # Font
                    if run.font.name:
                        run_styles.append(f'font-family: "{run.font.name}", serif')
                    
                    # Dimensione font
                    if run.font.size:
                        size_pt = run.font.size.pt
                        run_styles.append(f'font-size: {size_pt}pt')
                    
                    # Colore
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        color_hex = f'#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
                        run_styles.append(f'color: {color_hex}')
                    
                    # Bold, Italic, Underline
                    if run.bold:
                        run_styles.append('font-weight: bold')
                    if run.italic:
                        run_styles.append('font-style: italic')
                    if run.underline:
                        run_styles.append('text-decoration: underline')
                    
                    # Applica formattazione
                    if run_styles:
                        style_str = '; '.join(run_styles)
                        text = f'<span style="{style_str}">{text}</span>'
                    else:
                        # Usa tag semplici se non ci sono stili custom
                        if run.bold:
                            text = f'<strong>{text}</strong>'
                        if run.italic:
                            text = f'<em>{text}</em>'
                        if run.underline:
                            text = f'<u>{text}</u>'
                    
                    para_html += text
                
                para_html += f'</{tag}>'
                html_parts.append(para_html)
            
            # Gestisci tabelle (stile OnlyOffice)
            for table in doc.tables:
                html_parts.append('''
                    <table style="
                        border-collapse: collapse; 
                        width: 100%; 
                        margin: 12pt 0; 
                        border: 1px solid #000000;
                        font-size: 14pt;
                        font-family: 'Times New Roman', Times, serif;
                    ">
                ''')
                for row in table.rows:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Escape HTML
                        import html as html_module
                        cell_text = html_module.escape(cell_text)
                        html_parts.append(f'''
                            <td style="
                                padding: 4pt 8pt; 
                                border: 1px solid #000000;
                                vertical-align: top;
                            ">{cell_text}</td>
                        ''')
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            
            # Unisci tutto l'HTML
            html_content = '\n'.join(html_parts)
            
            result = {
                'html': html_content,
                'metadata': metadata,
                'success': True
            }
            
            # Cleanup file temporaneo se era una conversione .doc
            if temp_dir_to_cleanup and os.path.exists(temp_dir_to_cleanup):
                try:
                    shutil.rmtree(temp_dir_to_cleanup)
                    print(f"🧹 Pulizia file temporaneo .docx completata")
                except Exception as cleanup_error:
                    print(f"⚠️ Errore pulizia temp: {cleanup_error}")
            
            return result
            
        except Exception as e:
            # Cleanup anche in caso di errore
            if temp_dir_to_cleanup and os.path.exists(temp_dir_to_cleanup):
                try:
                    shutil.rmtree(temp_dir_to_cleanup)
                except:
                    pass
            
            return {
                'html': f'<p>❌ Errore conversione DOCX: {str(e)}</p>',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    @staticmethod
    def html_to_docx(html_content: str, output_path: str, metadata: Optional[Dict] = None) -> bool:
        """
        Converte HTML in DOCX
        
        Args:
            html_content: Contenuto HTML da convertire
            output_path: Percorso file DOCX di output
            metadata: Metadata opzionali (title, author, etc.)
            
        Returns:
            True se successo, False altrimenti
        """
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx non installato")
            return False
        
        try:
            doc = Document()
            
            # Imposta metadata
            if metadata:
                if 'title' in metadata:
                    doc.core_properties.title = metadata['title']
                if 'author' in metadata:
                    doc.core_properties.author = metadata['author']
                if 'subject' in metadata:
                    doc.core_properties.subject = metadata['subject']
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Converti elementi HTML in paragrafi DOCX
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div', 'table']):
                if element.name == 'table':
                    # Gestisci tabelle
                    rows = element.find_all('tr')
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                        table.style = 'Table Grid'
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                table.rows[i].cells[j].text = cell.get_text(strip=True)
                    
                else:
                    # Gestisci paragrafi
                    text = element.get_text(strip=True)
                    if not text:
                        continue
                    
                    # Crea paragrafo
                    para = doc.add_paragraph()
                    
                    # Imposta stile basato su tag
                    if element.name == 'h1':
                        para.style = 'Heading 1'
                    elif element.name == 'h2':
                        para.style = 'Heading 2'
                    elif element.name == 'h3':
                        para.style = 'Heading 3'
                    
                    # Aggiungi testo con formattazione
                    run = para.add_run(text)
                    
                    # Gestisci formattazione inline
                    if element.find('strong') or element.find('b'):
                        run.bold = True
                    if element.find('em') or element.find('i'):
                        run.italic = True
                    if element.find('u'):
                        run.underline = True
            
            # Salva documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Errore conversione HTML → DOCX: {e}")
            return False
    
    @staticmethod
    def extract_text_only(file_path: str) -> str:
        """
        Estrae solo il testo puro da un DOCX (senza formattazione)
        """
        if not DOCX_AVAILABLE:
            return "⚠️ python-docx non installato"
        
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            return f"❌ Errore: {e}"

